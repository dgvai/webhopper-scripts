# !pip install requests
# !pip install bs4
# !pip install emoji
# !pip install python-docx

import requests 
from bs4 import BeautifulSoup 
import json
import emoji
from collections import deque
from urllib.parse import urljoin
from docx import Document

class WebHopperCrawler:
    def __init__(self, url):
        self.domain = url
        self.paragraphs = {}
        self.visited = set()

    def paragraph_extractor(self, soup):
        """extract title and paragraphs from the page"""
        titleSoup = soup.find(id='title')
        textSoup = soup.find(id='text')
        
        title = titleSoup.get_text() if titleSoup else None
        text = textSoup.get_text() if textSoup else None
    
        if title in self.paragraphs:
            self.paragraphs[title].append(text)
        else:
            self.paragraphs[title] = [text]
        return title

    def DFS_crawler(self, url):
        """Depth First Search Algorithm"""
        if url in self.visited:
            return
        self.visited.add(url)
        
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        title = self.paragraph_extractor(soup)
        print(emoji.emojize("✅ crawling completed: "), title, url)
        
        links = soup.find_all('a', href=True)
        for link in links:
            next_url = f"{DOMAIN}{link.get('href')}"
            self.DFS_crawler(next_url)

    def BFS_crawler(self, url):
        """Breadth First Search Algorithm"""
        queue = deque([url])
        
        while queue:
            url = queue.popleft()
            if url in self.visited:
                continue
            self.visited.add(url)
            
            response = requests.get(url)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            title = self.paragraph_extractor(soup)
            print(emoji.emojize("✅ crawling completed: "), title, url)
    
            links = soup.find_all('a', href=True)
            for link in links:
                next_url = f"{DOMAIN}{link.get('href')}"
                if next_url not in self.visited:
                    queue.append(next_url)

    def document_generator(self, algoname):
        """Human readable document generator"""
        doc = Document()
        for i, title in enumerate(self.paragraphs):
            para = self.paragraphs[title][0]
            doc.add_heading(title)
            doc.add_paragraph(para)
        doc.save(f"webhopper_{algoname}.docx")

    def reset(self):
        """reset the state of the class"""
        self.paragraphs = {}
        self.visited = set()

    def execute(self):
        """execute the program step by step"""
        print("Starting webhopper...")
        print("Depth First Search (DFS)")
        self.DFS_crawler(self.domain)
        print("DFS Completed, Saving document...")
        self.document_generator("DFS")
        self.reset()
        print("Breadth First Search (BFS)")
        self.BFS_crawler(self.domain)
        print("BFS Completed, Saving document...")
        self.document_generator("BFS")
        print("Webhopper Completed!")

if __name__ == "__main__":
    # set the root domain to crawl
    DOMAIN = "https://webhopper-client.vercel.app"
    
    webhopper = WebHopperCrawler(url=DOMAIN)
    webhopper.execute()